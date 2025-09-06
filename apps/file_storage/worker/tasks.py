from celery import Celery
import PyPDF2
from docx import Document
import tempfile
from shared.storage.minio_client import minio_client
from shared.db.connection import AsyncSessionLocal
from ..infra.db.repositories import FileRepository
from config.settings import settings

celery_app = Celery(
    'file_storage_tasks',
    broker=settings.redis_url,
    backend=settings.redis_url
)


@celery_app.task
def extract_metadata(file_id: int):
    import asyncio
    return asyncio.run(_extract_metadata_async(file_id))


async def _extract_metadata_async(file_id: int):
    async with AsyncSessionLocal() as session:
        file_repo = FileRepository(session)

        file = await file_repo.get_by_id(file_id)
        if not file:
            return

        with tempfile.NamedTemporaryFile() as temp_file:
            try:
                data = minio_client.download_file(file.s3_path)
                temp_file.write(data.read())
                temp_file.flush()

                metadata = {}

                if file.content_type == "application/pdf":
                    metadata = _extract_pdf_metadata(temp_file.name)
                elif file.content_type in ["application/msword",
                                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                    metadata = _extract_docx_metadata(temp_file.name)

                await file_repo.update_metadata(file_id, metadata)
                await session.commit()

            except Exception as e:
                await file_repo.update_metadata(file_id, {"error": f"Failed to extract metadata: {e}"})
                await session.commit()


def _extract_pdf_metadata(file_path: str):
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            return {
                "pages": len(reader.pages),
                "author": reader.metadata.get("/Author", "Unknown") if reader.metadata else "Unknown",
                "title": reader.metadata.get("/Title", "Unknown") if reader.metadata else "Unknown",
                "creation_date": str(reader.metadata.get("/CreationDate", "Unknown")) if reader.metadata else "Unknown",
                "creator": reader.metadata.get("/Creator", "Unknown") if reader.metadata else "Unknown"
            }
    except Exception as e:
        return {"error": f"Could not extract PDF metadata: {e}"}


def _extract_docx_metadata(file_path: str):
    try:
        doc = Document(file_path)
        return {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "title": doc.core_properties.title or "Unknown",
            "author": doc.core_properties.author or "Unknown",
            "creation_date": str(doc.core_properties.created) if doc.core_properties.created else "Unknown"
        }
    except Exception as e:
        return {"error": f"Could not extract DOCX metadata: {e}"}